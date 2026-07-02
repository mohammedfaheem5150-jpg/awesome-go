"""Draft the monthly IR (ADB SAFEGATE AGL SPARE PARTS REQUISITION) as .docx.

Reads the damaged-fitting tracker workbook, takes one month's fittings
(by default only those with no IR NO yet), groups them by fitting type
and writes a requisition document in the team's real IR layout:

    header  : title / 'Requisition Type : Internal Purchase'
              'Date Requested : <date>    Requisition No: ADB-IR <n>'
    table   : S/N | ERP No. | DESCRIPTION | REQUIRED QTY | UNIT PRICE
              | TOTAL COST | REMARKS     (one row per fitting type)
    below   : JUSTIFICATION bullets + REQUESTED/ACKNOWLEDGED/APPROVED block

ERP numbers, descriptions and unit prices come from the config catalog
(`incident: ir: catalog:`) keyed by the fitting-type prefix (TEC,
LICATC, ...).  Types missing from the catalog still get a row, with
<FILL> placeholders, so the draft is always complete enough to finish
by hand.

Usage:
    python ir_draft.py 2025-11                # or NOV-2025
    python ir_draft.py 2025-11 --ir-no 55 --all   # include already-claimed rows
"""

import argparse
import datetime
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from incident_tracker import find_data_sheet, map_headers, _data_rows, _month_key
from openpyxl import load_workbook
from run_incident import HERE, load_config

DEFAULT_IR = {
    "requisition_type": "Internal Purchase",
    "requested_by": "",        # names typed on the final document
    "acknowledged_by": "",
    "approved_by": "",
    "currency": "AED",
    "justification": [
        "The light fixture was damaged due to unknown vehicle movement. "
        "These fixtures must be taken and utilized from ADA store. Cost will "
        "be under ADA-AGL as the assets are damaged due to the unknown "
        "vehicle hit.",
        "Incident reports for this IR has been attached.",
    ],
    "catalog": {},   # fitting-type prefix -> {erp_no, description, unit_price, remarks}
}


def _ordinal(d):
    return "%02d%s %s" % (d.day,
                          {1: "st", 2: "nd", 3: "rd"}.get(d.day % 10 if d.day not in (11, 12, 13) else 0, "th"),
                          d.strftime("%B %Y"))


def parse_month(text):
    text = text.strip().upper()
    for fmt in ("%Y-%m", "%b-%Y", "%B-%Y", "%m-%Y", "%Y%m"):
        try:
            return datetime.datetime.strptime(text, fmt)
        except ValueError:
            continue
    raise SystemExit("cannot parse month %r (use e.g. 2025-11 or NOV-2025)" % text)


def month_fittings(tracker_path, month_dt, include_claimed=False):
    wb = load_workbook(tracker_path)
    ws = find_data_sheet(wb)
    header_row, cmap = map_headers(ws)
    key = month_dt.strftime("%b-%Y").upper()
    rows = [r for r in _data_rows(ws, header_row, cmap) if _month_key(r["date"]) == key]
    if not include_claimed:
        rows = [r for r in rows if not r.get("ir_no")]
    return rows


def group_items(rows, catalog):
    groups = {}
    for r in rows:
        t = str(r.get("fitting_type") or "").strip() or "UNKNOWN"
        groups.setdefault(t, []).append(r)
    items = []
    for t in sorted(groups):
        cat = catalog.get(t, {})
        qty = len(groups[t])
        price = cat.get("unit_price")
        items.append({
            "fitting_type": t,
            "qty": qty,
            "erp_no": str(cat.get("erp_no") or "<FILL ERP No.>"),
            "description": str(cat.get("description")
                               or "<FILL DESCRIPTION for %s fittings>" % t),
            "unit_price": price,
            "total": round(qty * price, 2) if isinstance(price, (int, float)) else None,
            "remarks": str(cat.get("remarks") or ""),
            "fittings": groups[t],
        })
    return items


def _money(val, currency):
    if val is None:
        return "<FILL>"
    return ("{:,.2f}".format(val)).rstrip("0").rstrip(".") + " " + currency


def build_ir_docx(items, out_path, *, ir_no="<FILL>", date=None, cfg=None):
    cfg = {**DEFAULT_IR, **(cfg or {})}
    date = date or datetime.date.today()
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Pt(36)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ADB SAFEGATE AGL SPARE PARTS REQUISITION")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("Requisition Type\t:  %s" % cfg["requisition_type"])
    doc.add_paragraph("Date Requested\t:  %s\t\t\tRequisition No: ADB-IR %s"
                      % (_ordinal(date), ir_no))

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    for cell, head in zip(table.rows[0].cells,
                          ["S/N", "ERP No.", "DESCRIPTION", "REQUIRED QTY",
                           "UNIT PRICE", "TOTAL COST", "REMARKS"]):
        cell.text = head
        cell.paragraphs[0].runs[0].bold = True
    for i, item in enumerate(items, start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = item["erp_no"]
        row[2].text = item["description"]
        row[3].text = str(item["qty"])
        row[4].text = _money(item["unit_price"], cfg["currency"])
        row[5].text = _money(item["total"], cfg["currency"])
        row[6].text = item["remarks"]

    doc.add_paragraph("(NOT PAYABLE)").runs[0].bold = True
    doc.add_paragraph("JUSTIFICATION:").runs[0].bold = True
    for line in cfg["justification"]:
        doc.add_paragraph(line, style="List Bullet")

    # annex list: the per-fitting detail the incident reports back up
    doc.add_paragraph("ANNEX - DAMAGED FITTINGS COVERED BY THIS IR:").runs[0].bold = True
    annex = doc.add_table(rows=1, cols=5)
    annex.style = "Table Grid"
    for cell, head in zip(annex.rows[0].cells,
                          ["S/N", "DATE", "LOCATION", "FITTING REF", "INCIDENT REPORT / WO"]):
        cell.text = head
        cell.paragraphs[0].runs[0].bold = True
    n = 0
    for item in items:
        for r in sorted(item["fittings"], key=lambda x: (x["date"] or datetime.date.max)):
            n += 1
            row = annex.add_row().cells
            row[0].text = str(n)
            row[1].text = r["date"].strftime("%d-%b-%y") if r["date"] else ""
            row[2].text = str(r.get("location") or "")
            row[3].text = str(r.get("fitting_ref") or "")
            row[4].text = "%s / WO %s" % (r.get("report_id") or "?", r.get("wo_number") or "?")

    doc.add_paragraph("")
    sig = doc.add_table(rows=4, cols=3)
    sig.style = "Table Grid"
    heads = ["REQUESTED BY (ADB SAFEGATE)", "ACKNOWLEDGED BY (SINYAR)", "APPROVED BY (ADAC)"]
    names = [cfg["requested_by"], cfg["acknowledged_by"], cfg["approved_by"]]
    for c, head in enumerate(heads):
        cell = sig.rows[0].cells[c]
        cell.text = head
        cell.paragraphs[0].runs[0].bold = True
        sig.rows[1].cells[c].text = "Name: %s" % names[c]
        sig.rows[2].cells[c].text = "Signature:"
        sig.rows[3].cells[c].text = "Date:"

    doc.save(out_path)
    return out_path


def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("month", help="month to draft, e.g. 2025-11 or NOV-2025")
    ap.add_argument("--ir-no", default="<FILL>", help="requisition number (ADB-IR ...)")
    ap.add_argument("--all", action="store_true",
                    help="include fittings that already have an IR NO")
    ap.add_argument("--config", help="config yaml (default: config.yaml)")
    ap.add_argument("--out", help="output .docx path")
    args = ap.parse_args(argv)

    cfg = load_config(args.config)
    ir_cfg = {**DEFAULT_IR, **(cfg.get("ir") or {})}
    month_dt = parse_month(args.month)
    tracker_path = os.path.join(HERE, cfg["tracker_file"])
    if not os.path.exists(tracker_path):
        raise SystemExit("tracker workbook not found: %s (run run_incident.py first)"
                         % tracker_path)

    rows = month_fittings(tracker_path, month_dt, include_claimed=args.all)
    if not rows:
        raise SystemExit("no fittings pending IR for %s" % month_dt.strftime("%b-%Y").upper())
    items = group_items(rows, ir_cfg.get("catalog") or {})

    out = args.out or os.path.join(
        HERE, "ADB_IR_DRAFT_%s.docx" % month_dt.strftime("%b-%Y").upper())
    build_ir_docx(items, out, ir_no=args.ir_no, cfg=ir_cfg)
    missing = [i["fitting_type"] for i in items if i["unit_price"] is None]
    print("draft written: %s (%d line(s), %d fitting(s))"
          % (out, len(items), len(rows)))
    if missing:
        print("NOTE: no catalog entry for %s - <FILL> placeholders left in the draft; "
              "add them under incident: ir: catalog: in config.yaml"
              % ", ".join(missing))
    return 0


if __name__ == "__main__":
    sys.exit(main())
